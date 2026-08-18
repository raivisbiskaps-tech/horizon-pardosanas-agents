"""
app.py — Dokumentācijas Q&A aplikācija
=======================================
Streamlit web saskarne, kas ļauj uzdot jautājumus par dokumentāciju.

Lietošana:
    streamlit run app.py
"""

import os
import sys
import time
import io
import re
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
import streamlit as st
import pandas as pd
from dotenv import load_dotenv

load_dotenv()

# ── Konfigurācija (definē PIRMS funkcijām) ───────────────────────────────────

BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
CHROMA_DIR   = os.path.join(BASE_DIR, "chroma_db")
DOCS_DIR     = os.path.join(BASE_DIR, "docs")
COLLECTION_NAME = "dokumentacija"
TOP_K_RESULTS   = 30

# Pieejamie modeļi
MODELS = {
    "🟣 Claude Sonnet":  {"provider": "bedrock", "model": "anthropic.claude-sonnet-4-5"},
    "🟣 Claude Haiku":   {"provider": "bedrock", "model": "anthropic.claude-haiku-4-5-20251001"},
    "🟠 Mistral Small":  {"provider": "mistral", "model": "mistral-small-latest"},
    "🟠 Mistral Large":  {"provider": "mistral", "model": "mistral-large-2411"},
    "🔵 Gemini Flash":   {"provider": "gemini",  "model": "gemini-1.5-flash"},
    "🔵 Gemini Pro":     {"provider": "gemini",  "model": "gemini-1.5-pro"},
}

SYSTEM_PROMPT = """Tu esi pieredzējis ERP Horizon pārdošanas atbalsta speciālists. Strādā uzņēmumā VISMA, kas izplata, izstrādā un ievieš ERP Horizon.

Noteikumi:
- Balsties uz sev pieejamo informāciju un klienta sarakstē sniegtajiem faktiem
- Pilnīgas, detalizētas atbildes — ne pārāk īsas, ne pārāk garas
- Nekad neizdomā informāciju
- Ja jautājums ir neskaidrs, lūdz precizējumu
- Atbild tajā pašā valodā
- NEKAD neminē "dokumentus", "dokumentāciju" vai "avotus" — izmanto formulējumus: "pēc man pieejamās informācijas", "atbilstoši manām zināšanām", "cik man zināms" u.tml.
- Ja nevari atbildēt uz jautājumu → atbildi: "Atbilstoši manām zināšanām, uz šo jautājumu nevarēšu sniegt precīzu atbildi. Jūsu jautājumu esmu piefiksējis un nosūtīšu menedžerim, kurš ar jums sazināsies tuvākajā laikā." Pievieno marķieri [ACTION:NEZINA]

Konteksts par klientu:
- Potenciāls klients bez Horizon zināšanām
- Izvēlas ERP sistēmu — jānotur viņa interese
- Uzdod uzvedinošus jautājumus un piedāvā iegūt papildus informāciju
- Bez jautājumiem, klients var sarakstē iekļaut informāciju, kas nepieciešama piedāvājuma un līguma sagatavošanai — uzņēmuma rekvizīti, kontaktinformācija, nepieciešamā sistēmas komplektācija
- Ja kontekstā ir "[Automātiski iegūti rekvizīti no LR Uzņēmumu reģistra — apstipriniet pie klienta]" — nolasi tos un apstiprina pie klienta VISUS pieejamos laukus: nosaukumu, uzņēmuma tipu, reģistrācijas numuru, juridisko adresi, PVN reģistrācijas numuru un statusu. Formulē kā: "Atradu šādus datus par jūsu uzņēmumu — lūdzu apstipriniet, vai viss ir pareizi: [nosaukums, tips, reģ. nr., adrese, PVN nr. un statuss]"
- Ja klients sarunā piemin uzņēmuma nosaukumu un kontekstā nav rekvizītu — lūdz klientu apstiprināt uzņēmumu: "Vai pareizi sapratu — jūs pārstāvat [nosaukums]?"

Ziņojumu veidi — OBLIGĀTI ievēro:

1. JAUTĀJUMS par Horizon (piemēram: "Kādi moduļi ir pieejami?", "Cik maksā?")
   → Atbildi, balstoties uz sev pieejamo informāciju

2. ATBILDE uz mūsu jautājumu (piemēram: "Jā", "Nē", "Vajag arī personālu", "10 lietotāji", "Tieši tā")
   → Interpretē kā atbildi uz pēdējo uzdoto jautājumu. NEKAD nesaki "nevaru sniegt atbildi". Apkopo saprasto un turpini.

3. INFORMĀCIJAS sniegšana (piemēram: uzņēmuma nosaukums, reģ. nr., kontaktpersona, adrese)
   → Uztver kā faktu. Apstiprina saņemšanu ("Paldies, piefiksēju!"). NEKAD nemēģini pārbaudīt vai apstrīdēt šo informāciju.

4. UZDEVUMS vai KOMANDA (piemēram: "Sagatavo līgumu", "Gatavo tāmi", "Nosūti piedāvājumu")
   → Izpildi vai apstiprina nodomu. NEKAD nesaki "nevaru sniegt atbildi".

Ja nav skaidrs uz kuru jautājumu klients atbild — pārjautā konkrēti.

Terminoloģija:
- "Bizness" un "Ražošana" = Horizon papildiespēju paku nosaukumi (nevis vispārīgi vārdi)

Stils:
- Viegla, sarunbiedra valoda
- Var izmantot humoru un iepīt atbildēs pa kādam jokam

Tabulas:
- Izmantot Markdown tabulas, ja atbilde ir apjomīga strukturēta informācija (izmaksas, cenas u.tml.)
- Tabulas galvenei izmanto treknrakstu (|Modulis|Cena| u.tml.)
- Ciparus un cenas formatē konsekventi

Darbību marķieri — ieviešanas tāmes un piedāvājuma sagatavošanai:
- Izvērtē sarakstes kontekstu — ko jau zinām
- Uzdod tikai trūkstošos precizējošos jautājumus (moduļi, lietotāju skaits u.tml.)
- Apkopo saprasto un pārjautā klientam
- [ACTION:TAME] — pievieno pēc klienta apstiprinājuma par nepieciešamajiem moduļiem
- [ACTION:LIGUMS] — pievieno kad ir pietiekami daudz info līgumam; OBLIGĀTI jābūt zināmam uzņēmuma nosaukumam VAI reģistrācijas numuram, pretējā gadījumā vispirms papraси: "Lai sagatavotu līgumu, lūdzu norādiet uzņēmuma nosaukumu un reģistrācijas numuru."
- [ACTION:NEZINA] — pievieno kad nevari atbildēt uz jautājumu
- Abus var pievienot vienlaikus, piemēram: [ACTION:TAME][ACTION:LIGUMS]
- Marķieri ir tehniski tagi — klients tos neredz, tos apstrādā sistēma automātiski
- SVARĪGI: marķierus raksti TIEŠI tā, bez tulkošanas, maiņas vai Markdown formatēšanas (bez **)"""


# ── Čata marķieri ────────────────────────────────────────────────────────────

CHAT_MARKERS = {
    "[ACTION:TAME]":   "tāme",
    "[ACTION:LIGUMS]": "līgums",
    "[ACTION:NEZINA]": "nezina",
}

def parse_markers(text: str) -> tuple[str, list[str]]:
    """Izvelk darbību marķierus no AI atbildes.
    Atgriež (tīrs teksts bez marķieriem, atrasto marķieru saraksts).
    """
    found = []
    # Izmanto regex — izturīgi pret Markdown formatēšanu (**) un citām variācijām
    # Piemēri: [ACTION:TAME], [**ACTION:TAME**], [ACTION TAME], utt.
    tame_pat   = re.compile(r'\[\*{0,2}ACTION[:\s_-]*TAME\*{0,2}\]',   re.IGNORECASE)
    ligums_pat = re.compile(r'\[\*{0,2}ACTION[:\s_-]*LIGUMS?\*{0,2}\]', re.IGNORECASE)
    nezina_pat = re.compile(r'\[\*{0,2}ACTION[:\s_-]*NEZINA\*{0,2}\]',  re.IGNORECASE)
    if tame_pat.search(text):
        text = tame_pat.sub("", text)
        found.append("tāme")
    if ligums_pat.search(text):
        text = ligums_pat.sub("", text)
        found.append("līgums")
    if nezina_pat.search(text):
        text = nezina_pat.sub("", text)
        found.append("nezina")
    return text.strip(), found


# ── Indeksēšana ───────────────────────────────────────────────────────────────

@st.cache_resource(show_spinner="Indeksē dokumentāciju, lūdzu uzgaidi...")
def auto_ingest():
    """Indeksē dokumentus pie katras jaunas sesijas."""
    sys.path.insert(0, BASE_DIR)
    from ingest import ingest
    ingest(docs_dir=DOCS_DIR)
    return True


# ── ChromaDB ──────────────────────────────────────────────────────────────────

@st.cache_resource(show_spinner="Ielādē dokumentu datubāzi...")
def load_collection():
    """Ielādē ChromaDB kolekciju."""
    import chromadb
    from chromadb.utils import embedding_functions

    if not os.path.exists(CHROMA_DIR):
        st.error("❌ Dokumentu datubāze nav atrasta. Pārliecinies, ka docs/ mapē ir dokumenti.")
        st.stop()

    client = chromadb.PersistentClient(path=CHROMA_DIR)
    emb_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="paraphrase-multilingual-mpnet-base-v2"
    )
    try:
        return client.get_collection(name=COLLECTION_NAME, embedding_function=emb_fn)
    except Exception as e:
        st.error(f"❌ Kolekcija nav atrasta: {e}")
        st.stop()


# ── Bedrock (Claude) klients ──────────────────────────────────────────────────

@st.cache_resource(show_spinner="Savieno ar Claude...")
def load_bedrock_client():
    """Inicializē Claude klientu caur Bedrock API key."""
    import anthropic
    api_key = os.getenv("BEDROCK_API_KEY") or st.secrets.get("BEDROCK_API_KEY", "")
    if not api_key:
        return None
    # Bedrock API key (ABSK...) — izmantojam kā parasto Anthropic klientu
    return anthropic.Anthropic(api_key=api_key)


# ── Mistral klients ───────────────────────────────────────────────────────────

@st.cache_resource(show_spinner="Savieno ar AI...")
def load_mistral_client():
    """Inicializē Mistral klientu."""
    from openai import OpenAI
    api_key = os.getenv("MISTRAL_API_KEY")
    if not api_key:
        return None
    return OpenAI(api_key=api_key, base_url="https://api.mistral.ai/v1")

@st.cache_resource(show_spinner="Savieno ar Gemini...")
def load_gemini_client():
    """Inicializē Gemini klientu caur OpenAI-saderīgo API."""
    from openai import OpenAI
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        return None
    return OpenAI(
        api_key=api_key,
        base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
    )


# ── Balss transkripcija ───────────────────────────────────────────────────────

# ── Balss transkripcija ───────────────────────────────────────────────────────

def transcribe_audio(audio_bytes: bytes, mime_type: str = "audio/webm") -> str | None:
    try:
        from groq import Groq
        api_key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY", "")
        if not api_key:
            st.warning("⚠️ GROQ_API_KEY nav iestatīts — balss ievade nav pieejama.")
            return None
        ext = mime_type.split("/")[-1].split(";")[0].strip()
        client = Groq(api_key=api_key)
        transcription = client.audio.transcriptions.create(
            file=(f"audio.{ext}", audio_bytes, mime_type.split(";")[0]),
            model="whisper-large-v3",
            language="lv",
        )
        return transcription.text.strip()
    except Exception as e:
        st.error(f"❌ Balss atpazīšanas kļūda: {e}")
        return None


# ── Mic pogas komponents ──────────────────────────────────────────────────────

@st.cache_resource
def _get_mic_button_fn():
    import streamlit.components.v1 as _stc
    path = os.path.join(BASE_DIR, "components", "mic_button")
    return _stc.declare_component("mic_button", path=path)


def mic_button(reset_counter: int = 0) -> dict | None:
    fn = _get_mic_button_fn()
    return fn(reset_counter=reset_counter, key="mic_button", default=None)


# ── RAG ───────────────────────────────────────────────────────────────────────

def retrieve_context(collection, question: str) -> tuple[str, list[str]]:
    """Meklē relevantos fragmentus ChromaDB."""
    results = collection.query(
        query_texts=[question],
        n_results=TOP_K_RESULTS,
        include=["documents", "metadatas", "distances"],
    )
    context_parts, sources = [], []
    for doc, meta, dist in zip(
        results["documents"][0],
        results["metadatas"][0],
        results["distances"][0],
    ):
        if dist < 0.95:
            src = meta.get("source", "nezināms")
            context_parts.append(f"[Avots: {src}]\n{doc}")
            if src not in sources:
                sources.append(src)
    return "\n\n---\n\n".join(context_parts), sources


def ask_ai(question: str, context: str, model_name: str,
           history: list = None, izmanto_rag: bool = True) -> str:
    """Nosūta jautājumu uz izvēlēto AI modeli, iekļaujot sarakstes vēsturi.

    izmanto_rag=True  → ziņojums ietver dokumentu fragmentus un to ierobežojumu
    izmanto_rag=False → ziņojums ir tikai sarakstes turpinājums bez RAG ierobežojuma
    """
    model_cfg = MODELS[model_name]
    provider  = model_cfg["provider"]
    model_id  = model_cfg["model"]

    if izmanto_rag and context:
        user_message = f"""Zemāk ir pieejamā informācija, ko drīksti izmantot atbildē.
Ja atbilde nav šajā informācijā — nekādā gadījumā to neizdomā. NEKAD neatsaucies uz "fragmentiem", "dokumentiem" vai "dokumentāciju" — izmanto: "pēc man pieejamās informācijas", "atbilstoši manām zināšanām".

=== PIEEJAMĀ INFORMĀCIJA ===
{context}
=== BEIGAS ===

Jautājums: {question}"""
    else:
        # Atbilde/komanda/informācija — turpina sarakstes loģiku bez RAG ierobežojuma
        user_message = question

    # Veido ziņojumu sarakstu ar sarakstes vēsturi
    messages = [{"role": "system", "content": SYSTEM_PROMPT}]

    # Pievieno iepriekšējos ziņojumus (bez pēdējā user ziņojuma, kas jau ir user_message)
    if history:
        for msg in history[:-1]:
            messages.append({
                "role": msg["role"],
                "content": msg["content"],
            })

    # Pievieno pašreizējo jautājumu ar dokumentu kontekstu
    messages.append({"role": "user", "content": user_message})

    max_retries = 3
    for attempt in range(max_retries):
        try:
            if provider == "bedrock":
                client = load_bedrock_client()
                if not client:
                    return "❌ BEDROCK_API_KEY nav iestatīts Streamlit Secrets."
                # Anthropic API: system ir atsevišķs parametrs, nevis messages sarakstā
                bedrock_messages = [m for m in messages if m["role"] != "system"]
                response = client.messages.create(
                    model=model_id,
                    max_tokens=2048,
                    system=SYSTEM_PROMPT,
                    messages=bedrock_messages,
                )
                return response.content[0].text

            elif provider == "mistral":
                client = load_mistral_client()
                if not client:
                    return "❌ MISTRAL_API_KEY nav iestatīts Streamlit Secrets."
                response = client.chat.completions.create(
                    model=model_id,
                    messages=messages,
                )
                return response.choices[0].message.content

            elif provider == "gemini":
                client = load_gemini_client()
                if not client:
                    return "❌ GOOGLE_API_KEY nav iestatīts Streamlit Secrets."
                response = client.chat.completions.create(
                    model=model_id,
                    messages=messages,
                )
                return response.choices[0].message.content

        except Exception as e:
            error_str = str(e).lower()
            if "rate" in error_str or "limit" in error_str or "429" in error_str:
                if attempt < max_retries - 1:
                    time.sleep(5 * (attempt + 1))
                    continue
                return "⚠️ Sistēma šobrīd ir noslogota. Lūdzu, mēģini vēlreiz pēc dažām sekundēm."
            return f"❌ Kļūda: {str(e)}"


# ── Qwilr integrācija ────────────────────────────────────────────────────────

def summarize_chat_for_proposal(messages: list, model_name: str) -> dict:
    """Izmanto AI, lai no sarakstes iegūtu strukturētu kopsavilkumu piedāvājumam."""
    if not messages:
        return {}

    history_text = ""
    for msg in messages:
        role = "Klients" if msg["role"] == "user" else "Aģents"
        history_text += f"{role}: {msg['content']}\n\n"

    prompt = f"""No šīs sarakstes iegūsti strukturētu informāciju piedāvājumam. Atbildi TIKAI JSON formātā, bez papildu teksta.

Sarakstes vēsture:
{history_text}

Atbildi šādā JSON formātā:
{{
  "klients": "Klienta vārds vai uzņēmuma nosaukums, ja tas minēts sarakstē. Ja nav minēts — 'Nav norādīts'",
  "client_interests": "Īss kopsavilkums par ko klients interesējas (1-2 teikumi)",
  "modules": "Pieminētie Horizon moduļi vai pakotnes (vai 'Nav precizēts')",
  "key_questions": "Galvenie klienta jautājumi (1-3 punkti)",
  "next_steps": "Ieteicamie nākamie soļi (1-2 teikumi)"
}}"""

    model_cfg = MODELS[model_name]
    try:
        if model_cfg["provider"] == "mistral":
            client = load_mistral_client()
            if not client:
                return {}
        else:
            client = load_gemini_client()
            if not client:
                return {}

        response = client.chat.completions.create(
            model=model_cfg["model"],
            messages=[{"role": "user", "content": prompt}],
        )
        import json
        text = response.choices[0].message.content.strip()
        # Izņem JSON no atbildes
        if "```" in text:
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
        return json.loads(text)
    except Exception:
        return {}


def create_qwilr_proposal(messages: list, model_name: str) -> tuple[bool, str]:
    """Nosūta sarakstes datus uz Zapier Webhook, kas izveido Qwilr piedāvājumu."""
    webhook_url = os.getenv("ZAPIER_WEBHOOK_URL")

    if not webhook_url:
        return False, "❌ ZAPIER_WEBHOOK_URL nav iestatīts Streamlit Secrets."
    if not messages:
        return False, "❌ Čats ir tukšs — nav ko iekļaut piedāvājumā."

    # Iegūst sarakstes kopsavilkumu
    summary = summarize_chat_for_proposal(messages, model_name)
    timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")

    # Sagatavo datus Zapier webhook
    payload = {
        "Klients":           summary.get("klients", "Nav norādīts"),
        "klienta_intereses": summary.get("client_interests", "Nav norādīts"),
        "moduli":            summary.get("modules", "Nav precizēts"),
        "galvenie_jautajumi":summary.get("key_questions", "Nav norādīts"),
        "nakamie_soli":      summary.get("next_steps", "Nav norādīts"),
        "datums":            timestamp,
        "nosaukums":         f"Horizon piedāvājums — {timestamp}",
    }

    import requests as req
    try:
        response = req.post(webhook_url, json=payload, timeout=15)
        if not response.ok:
            return False, f"❌ Zapier kļūda {response.status_code}: {response.text}"
        return True, "✅ Dati nosūtīti uz Zapier — Qwilr piedāvājums tiek gatavots!"
    except Exception as e:
        return False, f"❌ Kļūda: {str(e)}"


# ── Excel eksports ────────────────────────────────────────────────────────────

def extract_markdown_tables(text: str) -> list[pd.DataFrame]:
    """Parsē visas Markdown tabulas no teksta un atgriež DataFrame sarakstu."""
    tables = []
    # Atrod tabulas blokus (vismaz divas rindas ar | simbolu)
    table_pattern = re.compile(
        r'(\|.+\|\n\|[-| :]+\|\n(?:\|.+\|\n?)+)',
        re.MULTILINE
    )
    for match in table_pattern.finditer(text):
        lines = [l.strip() for l in match.group(0).strip().splitlines() if l.strip()]
        # Izfiltrē atdalītāja rindu (---|---|---)
        data_lines = [l for l in lines if not re.match(r'^\|[-| :]+\|$', l)]
        rows = []
        for line in data_lines:
            # Sadala pēc | un notīra atstarpes un treknrakstu **...**
            cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip())
                     for c in line.strip('|').split('|')]
            rows.append(cells)
        if len(rows) >= 2:
            df = pd.DataFrame(rows[1:], columns=rows[0])
            tables.append(df)
    return tables


def tables_to_excel_bytes(tables: list[pd.DataFrame]) -> bytes:
    """Pārvērš DataFrame sarakstu Excel faila baitiem."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for i, df in enumerate(tables):
            sheet_name = f"Tabula_{i+1}" if len(tables) > 1 else "Dati"
            df.to_excel(writer, index=False, sheet_name=sheet_name)
    return buf.getvalue()


# ── E-pasta sūtīšana ─────────────────────────────────────────────────────────

def _email_recipients(user_email: str = "") -> list[str]:
    """Atgriež saņēmēju sarakstu: TARGET_EMAIL + pieslēgtais lietotājs (ja atšķiras)."""
    target = os.getenv("TARGET_EMAIL", "")
    recipients = [target] if target else []
    if user_email and user_email.lower() not in [r.lower() for r in recipients]:
        recipients.append(user_email)
    return recipients


def send_file_by_email(
    file_bytes: bytes,
    filename: str,
    subject: str,
    body: str,
    user_email: str = "",
) -> tuple[bool, str]:
    """Nosūta failu kā e-pasta pielikumu uz TARGET_EMAIL un pieslēgtā lietotāja e-pastu."""
    from email.mime.base import MIMEBase
    from email import encoders

    gmail_user     = os.getenv("GMAIL_USER")
    gmail_password = os.getenv("GMAIL_APP_PASSWORD")
    recipients     = _email_recipients(user_email)

    if not all([gmail_user, gmail_password]) or not recipients:
        return False, "❌ Nav iestatīti e-pasta mainīgie."

    email_msg = MIMEMultipart()
    email_msg["From"]    = gmail_user
    email_msg["To"]      = ", ".join(recipients)
    email_msg["Subject"] = subject
    email_msg.attach(MIMEText(body, "plain", "utf-8"))

    # Pielikums
    part = MIMEBase("application", "octet-stream")
    part.set_payload(file_bytes)
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
    email_msg.attach(part)

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(gmail_user, gmail_password)
            server.sendmail(gmail_user, recipients, email_msg.as_string())
        return True, ", ".join(recipients)
    except Exception as e:
        return False, str(e)

def send_chat_by_email(messages: list, user_email: str = "") -> tuple[bool, str]:
    """Nosūta čata vēsturi uz TARGET_EMAIL un pieslēgtā lietotāja e-pastu."""
    gmail_user     = os.getenv("GMAIL_USER")
    gmail_password = os.getenv("GMAIL_APP_PASSWORD")
    target_email   = os.getenv("TARGET_EMAIL")

    recipients = _email_recipients(user_email)

    if not all([gmail_user, gmail_password]) or not recipients:
        return False, "❌ Nav iestatīti e-pasta mainīgie (GMAIL_USER, GMAIL_APP_PASSWORD, TARGET_EMAIL)."

    if not messages:
        return False, "❌ Čats ir tukšs — nav ko sūtīt."

    # Veido e-pasta saturu
    timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
    body = f"Horizon pārdošanas aģents — saruna {timestamp}\n"
    if user_email:
        body += f"Lietotājs: {user_email}\n"
    body += "=" * 60 + "\n\n"

    for msg in messages:
        role  = "👤 Klients" if msg["role"] == "user" else "🤖 Aģents"
        body += f"{role}:\n{msg['content']}\n\n"
        body += "-" * 40 + "\n\n"

    # Veido e-pastu
    email_msg = MIMEMultipart()
    email_msg["From"]    = gmail_user
    email_msg["To"]      = ", ".join(recipients)
    email_msg["Subject"] = f"Horizon aģents — saruna {timestamp}"
    email_msg.attach(MIMEText(body, "plain", "utf-8"))

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(gmail_user, gmail_password)
            server.sendmail(gmail_user, recipients, email_msg.as_string())
        saņēmēji_teksts = " un ".join(recipients)
        return True, f"✅ Saruna nosūtīta uz {saņēmēji_teksts}"
    except Exception as e:
        return False, f"❌ Sūtīšanas kļūda: {e}"


def send_unanswered_question(question: str, user_email: str = "") -> tuple[bool, str]:
    """Nosūta neatbildēto klienta jautājumu uz menedžera e-pastu."""
    gmail_user     = os.getenv("GMAIL_USER")
    gmail_password = os.getenv("GMAIL_APP_PASSWORD")
    recipients     = _email_recipients(user_email)

    if not all([gmail_user, gmail_password]) or not recipients:
        return False, ""

    timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
    body = (
        f"Horizon pārdošanas aģents — neatbildēts jautājums\n"
        f"Datums: {timestamp}\n"
        f"Lietotājs: {user_email or '—'}\n"
        f"{'=' * 50}\n\n"
        f"{question}\n"
    )
    email_msg = MIMEMultipart()
    email_msg["From"]    = gmail_user
    email_msg["To"]      = ", ".join(recipients)
    email_msg["Subject"] = f"❓ Neatbildēts jautājums — {timestamp}"
    email_msg.attach(MIMEText(body, "plain", "utf-8"))

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(gmail_user, gmail_password)
            server.sendmail(gmail_user, recipients, email_msg.as_string())
        return True, ", ".join(recipients)
    except Exception as e:
        return False, str(e)


# ── Uzņēmumu reģistrs caur data.gov.lv API ───────────────────────────────────

_UR_RESOURCE_ID  = "25e80bf3-f107-4ab4-89ef-251b5b9374e9"
_VID_RESOURCE_ID = "610910e9-e086-4c5b-a7ea-0a896a697672"
_DGL_API_URL     = "https://data.gov.lv/api/3/action/datastore_search"


def fetch_pvn_status(regcode_str: str) -> dict:
    """Pārbauda PVN maksātāja statusu VID datubāzē pēc reģistrācijas numura.

    Atgriež:
      pvn_numurs   – "LVxxxxxxxxxx" vai ""
      pvn_aktīvs   – True/False
      pvn_reģ      – reģistrācijas datums kā PVN maksātājs
      pvn_izsl     – izslēgšanas datums (ja bijis un izslēgts)
    """
    import requests as req
    import json as _json

    pvn_num = f"LV{regcode_str}"
    try:
        r = req.get(
            _DGL_API_URL,
            params={
                "resource_id": _VID_RESOURCE_ID,
                "filters":     _json.dumps({"Numurs": pvn_num}),
                "limit":       1,
            },
            timeout=8,
        )
        r.raise_for_status()
        records = r.json().get("result", {}).get("records", [])
    except Exception:
        return {"pvn_numurs": "", "pvn_aktīvs": None, "pvn_reģ": "", "pvn_izsl": ""}

    if not records:
        return {"pvn_numurs": "", "pvn_aktīvs": False, "pvn_reģ": "", "pvn_izsl": ""}

    rec = records[0]
    aktīvs = rec.get("Aktivs", "").strip().lower() == "ir"
    return {
        "pvn_numurs": pvn_num if aktīvs or rec.get("Izslegts", "").strip() else "",
        "pvn_aktīvs": aktīvs,
        "pvn_reģ":    rec.get("Registrets", "").strip(),
        "pvn_izsl":   rec.get("Izslegts", "").strip(),
    }


def fetch_data_gov_lv(vaicajums: str) -> dict:
    """Iegūst uzņēmuma rekvizītus no LR Uzņēmumu reģistra caur data.gov.lv API.

    Pieejamie lauki: nosaukums, reģ. numurs, juridiskā adrese, uzņēmuma tips,
    reģistrācijas datums, SEPA identifikators.
    """
    import requests as req
    import json as _json

    vaicajums = vaicajums.strip()

    # PVN numuru normalizē → 11 cipari
    if re.match(r'^LV\d{11}$', vaicajums, re.IGNORECASE):
        vaicajums = vaicajums[2:]

    is_regnum = bool(re.match(r'^\d{11}$', vaicajums))

    if is_regnum:
        params = {
            "resource_id": _UR_RESOURCE_ID,
            "filters":     _json.dumps({"regcode": int(vaicajums)}),
            "limit":       1,
        }
    else:
        params = {
            "resource_id": _UR_RESOURCE_ID,
            "q":           vaicajums,
            "limit":       15,
        }

    try:
        r = req.get(_UR_API_URL, params=params, timeout=10)
        r.raise_for_status()
        data = r.json()
    except Exception as e:
        return {"kļūda": f"data.gov.lv nav pieejams: {e}"}

    if not data.get("success"):
        return {"kļūda": "data.gov.lv API kļūda"}

    records = data.get("result", {}).get("records", [])
    if not records:
        return {"kļūda": f"Uzņēmums '{vaicajums}' nav atrasts Uzņēmumu reģistrā"}

    # Priekšroka aktīvajiem uzņēmumiem (terminated IS NULL)
    active = [rec for rec in records if not rec.get("terminated")]
    rec = active[0] if active else records[0]

    regcode_str = str(int(rec["regcode"])) if rec.get("regcode") else ""

    # PVN statuss no VID
    pvn = fetch_pvn_status(regcode_str) if regcode_str else {}

    return {
        "nosaukums":        rec.get("name", ""),
        "reg_numurs":       regcode_str,
        "juridiska_adrese": rec.get("address", ""),
        "tips":             rec.get("type_text", ""),
        "registrēts":       (rec.get("registered") or "")[:10],
        "sepa":             rec.get("sepa") or "",
        "pvn_numurs":       pvn.get("pvn_numurs", ""),
        "pvn_aktīvs":       pvn.get("pvn_aktīvs"),   # True / False / None
        "pvn_reģ":          pvn.get("pvn_reģ", ""),
        "pvn_izsl":         pvn.get("pvn_izsl", ""),
        "url":              f"https://ur.gov.lv/lv/meklesana/?term={regcode_str}",
    }


# ── Līguma sagatavošana ───────────────────────────────────────────────────────

# Šie mainīgo nosaukumi jālieto Word šablonā: {{ mainīgā_nosaukums }}
LIGUMA_MAINĪGIE_NOKLUSĒJUMS = {
    # Klienta rekvizīti (daļēji no firmas.lv)
    "uznemuma_nosaukums":    "",   # {{ uznemuma_nosaukums }}
    "reg_numurs":            "",   # {{ reg_numurs }}
    "pvn_numurs":            "",   # {{ pvn_numurs }}
    "juridiska_adrese":      "",   # {{ juridiska_adrese }}
    "fakt_adrese":           "",   # {{ fakt_adrese }}
    "talrunis":              "",   # {{ talrunis }}
    "epasts":                "",   # {{ epasts }}
    # Parakstītājs (no sarakstes)
    "paraksta":              "",   # {{ paraksta }}
    "paraksta_amats":        "",   # {{ paraksta_amats }}
    "pamat_uz":              "",   # {{ pamat_uz }} — statūti / prokūra / pilnvara
    # Kontaktpersona (no sarakstes)
    "kontaktpersona":        "",   # {{ kontaktpersona }}
    "kontaktpersonas_amats": "",   # {{ kontaktpersonas_amats }}
    "kontaktp_epasts":       "",   # {{ kontaktp_epasts }}
    "kontaktp_talrunis":     "",   # {{ kontaktp_talrunis }}
    # Līguma dati
    "datums":                "",   # {{ datums }} — aizpildās automātiski
    "liguma_numurs":         "",   # {{ liguma_numurs }} — jāaizpilda manuāli
    "lpp":                   "",   # {{ lpp }} — lappušu skaits, jāaizpilda manuāli
}


def extract_liguma_mainīgie(messages: list, model_name: str, rekviziti: dict = None) -> dict:
    """Izvelk līguma aizpildīšanai nepieciešamos mainīgos no sarakstes un rekvizītiem."""
    import json as _json

    mainīgie = dict(LIGUMA_MAINĪGIE_NOKLUSĒJUMS)
    mainīgie["datums"] = datetime.now().strftime("%d.%m.%Y")

    # 1. Aizpilda no firmas.lv rekvizītiem (ja ir)
    if rekviziti:
        mainīgie.update({
            "uznemuma_nosaukums": rekviziti.get("nosaukums", ""),
            "reg_numurs":         rekviziti.get("reg_numurs", ""),
            "pvn_numurs":         rekviziti.get("pvn_numurs", ""),
            "juridiska_adrese":   rekviziti.get("juridiska_adrese", ""),
            "talrunis":           rekviziti.get("talrunis", ""),
            "epasts":             rekviziti.get("epasts", ""),
        })

    if not messages:
        return mainīgie

    # 2. AI izvelk trūkstošo info no sarakstes
    history_text = "\n".join([
        f"{'Klients' if m['role'] == 'user' else 'Aģents'}: {m['content']}"
        for m in messages
    ])

    jau_zinami = {k: v for k, v in mainīgie.items() if v}
    jau_zinami_teksts = "\n".join(f"  {k}: {v}" for k, v in jau_zinami.items()) or "  (nekas nav zināms)"

    prompt = f"""No šīs pārdošanas sarakstes izvelc informāciju līgumam. Atbildi TIKAI ar JSON objektu, bez papildu teksta.

Sarakstes vēsture:
{history_text[-4000:]}

Jau zināmie dati (neaizstāj ar tukšiem!):
{jau_zinami_teksts}

Izvelc šādus laukus (ja nav atrodams — atstāj ""):
{{
  "uznemuma_nosaukums": "uzņēmuma pilnais nosaukums",
  "paraksta": "personas vārds, uzvārds, kas parakstīs līgumu",
  "paraksta_amats": "parakstītāja amats",
  "pamat_uz": "paraksta pamatojoties uz (piemēram: statūtiem, prokūru, pilnvaru Nr. X)",
  "fakt_adrese": "faktiskā adrese, ja atšķiras no juridiskās",
  "kontaktpersona": "kontaktpersonas vārds, uzvārds",
  "kontaktpersonas_amats": "kontaktpersonas amats",
  "kontaktp_epasts": "kontaktpersonas e-pasts",
  "kontaktp_talrunis": "kontaktpersonas tālrunis",
  "reg_numurs": "reģistrācijas numurs, ja minēts"
}}"""

    model_cfg = MODELS[model_name]
    try:
        if model_cfg["provider"] == "mistral":
            client = load_mistral_client()
        else:
            client = load_gemini_client()
        if not client:
            return mainīgie

        response = client.chat.completions.create(
            model=model_cfg["model"],
            messages=[{"role": "user", "content": prompt}],
        )
        text = response.choices[0].message.content.strip()
        # Izņem JSON no atbildes (arī ja ietīts ```json ... ```)
        json_match = re.search(r'\{[^{}]*\}', text, re.DOTALL)
        if json_match:
            extracted = _json.loads(json_match.group(0))
            for k, v in extracted.items():
                if v and k in mainīgie and not mainīgie.get(k):
                    mainīgie[k] = str(v)
    except Exception:
        pass

    return mainīgie


_TUKSS_PREFIKS = "_TUKSS_"  # + lauka nosaukums, piemēram "_TUKSS_uznemuma_nosaukums"
_LPP           = "_LPP_"   # marķieris lpp → NUMPAGES lauka kods


def _postprocess_ligums(doc_bytes: bytes) -> bytes:
    """Pēc docxtpl renderēšanas:
    - aizstāj _AIZPILDIT_ ar dzeltenā iezīmētām atstarpēm
    - aizstāj _LPP_ ar Word NUMPAGES lauka kodu
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_COLOR_INDEX

    buf = io.BytesIO(doc_bytes)
    doc = Document(buf)

    def _numpages_xml(run):
        """Iestata NUMPAGES lauku dotajā run elementā."""
        run.text = ""
        fc_begin = OxmlElement("w:fldChar")
        fc_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " NUMPAGES "
        fc_sep = OxmlElement("w:fldChar")
        fc_sep.set(qn("w:fldCharType"), "separate")
        fc_end = OxmlElement("w:fldChar")
        fc_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fc_begin, instr, fc_sep, fc_end])

    def _process_run(run):
        if _LPP in run.text:
            # Aizstāj ar NUMPAGES lauku
            prefix, suffix = run.text.split(_LPP, 1)
            run.text = prefix
            _numpages_xml(run)
            # Ja pēc marķiera ir teksts — pievieno jaunu run
            if suffix:
                new_run = OxmlElement("w:r")
                new_t = OxmlElement("w:t")
                new_t.text = suffix
                new_run.append(new_t)
                run._r.addnext(new_run)
        elif _TUKSS_PREFIKS in run.text:
            # Izvelk lauka nosaukumu no marķiera un iezīmē dzeltenā
            # Piemēram "_TUKSS_uznemuma_nosaukums" → "uznemuma_nosaukums" (dzeltenā)
            run.text = run.text.replace(_TUKSS_PREFIKS, "")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for para in doc.paragraphs:
        for run in para.runs:
            _process_run(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _process_run(run)

    buf_out = io.BytesIO()
    doc.save(buf_out)
    return buf_out.getvalue()


def generate_ligums_docx(messages: list, model_name: str,
                          rekviziti: dict = None) -> tuple:
    """Ģenerē līguma Word dokumentu no docxtpl šablona.

    Atgriež: (bytes, mainīgie_dict) ja veiksmīgi, (None, kļūdas_teksts) ja neizdevās.
    """
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return None, "❌ docxtpl nav instalēts. Pievieno 'docxtpl' requirements.txt un pārinstallē."

    template_path = os.path.join(BASE_DIR, "assets", "liguma_sablons.docx")
    if not os.path.exists(template_path):
        return None, (
            "❌ Šablons nav atrasts (assets/liguma_sablons.docx).\n"
            "Izveido Word dokumentu ar {{ mainīgais }} atzīmēm un saglabā kā assets/liguma_sablons.docx."
        )

    # Izvelk mainīgos no sarakstes + rekvizītiem
    mainīgie = extract_liguma_mainīgie(messages, model_name, rekviziti)

    # Sagatavo kontekstu docxtpl: tukšie → marķieris, lpp → LPP marķieris
    konteksts = {}
    for k, v in mainīgie.items():
        if k == "lpp":
            konteksts[k] = _LPP          # aizstāj ar NUMPAGES pēc renderēšanas
        elif not v:
            konteksts[k] = f"{_TUKSS_PREFIKS}{k}"  # "_TUKSS_uznemuma_nosaukums" u.tml.
        else:
            konteksts[k] = v

    try:
        doc = DocxTemplate(template_path)
        doc.render(konteksts)
        buf = io.BytesIO()
        doc.save(buf)
        # Pēcapstrāde: dzeltenā iezīmēšana + NUMPAGES lauks
        doc_bytes = _postprocess_ligums(buf.getvalue())
        return doc_bytes, mainīgie
    except Exception as e:
        return None, f"❌ Kļūda aizpildot šablonu: {e}"


# ── Ieviešanas tāme ───────────────────────────────────────────────────────────

# Visi iespējamie bloki — atslēgvārdi no sarakstes → atbilstošās sadaļas šablonā
TAME_BLOCKS = {
    "instalācija":              ["Instalācija"],
    "pamatsistēma":             ["Instalācija", "Pamatsistēma"],
    "grāmatvedība":             ["Instalācija", "Pamatsistēma"],
    "algas":                    ["Instalācija", "Pamatsistēma"],
    "personāls":                ["Instalācija", "Pamatsistēma", "Personāls"],
    "hop personāls":            ["Instalācija", "Pamatsistēma", "Personāls"],
    "pieteikumi":               ["Personāls"],
    "komandējumi":              ["Personāls"],
    "mani izdevumi":            ["Personāls"],
    "rīkojumi":                 ["Personāls"],
    "darba laika uzskaite":     ["Personāls"],
    "hop rēķini":               ["Grāmatvedība +"],
    "rēķini":                   ["Grāmatvedība +"],
    "numo":                     ["Numo"],
    "darba laika plānošana":    ["Numo"],
}

TAME_CLARIFYING_QUESTIONS = """Ja klients jautā par ieviešanu vai ieviešanas izmaksām, OBLIGĀTI uzdod šos precizējošos jautājumus PIRMS tāmes sagatavošanas:
1. Vai nepieciešama Pamatsistēma (grāmatvedība, rēķini, noliktava)?
2. Vai nepieciešams Algu un personāla modulis?
3. Vai nepieciešams HOP Personāls (darbinieku pieteikumi, komandējumi, izdevumi, rīkojumi)?
4. Vai nepieciešama HOP Darba laika uzskaite?
5. Vai nepieciešami HOP Rēķini (rēķinu saskaņošana)?
6. Vai nepieciešama NUMO Darba laika plānošana?
7. Cik lietotāji strādās sistēmā?
Neģenerē tāmi kamēr nav saņemtas atbildes uz šiem jautājumiem."""


def determine_tame_sections(messages: list, model_name: str) -> list[str]:
    """AI nosaka kuras sadaļas iekļaut tāmē, balstoties uz sarakstes."""
    history_text = ""
    for msg in messages:
        role = "Klients" if msg["role"] == "user" else "Aģents"
        history_text += f"{role}: {msg['content']}\n\n"

    prompt = f"""No šīs sarakstes nosaki, kuras Horizon ieviešanas sadaļas ir nepieciešamas klientam.
Atbildi TIKAI ar JSON sarakstu no šiem variantiem (iekļauj tikai vajadzīgos):
["Instalācija", "Pamatsistēma", "Personāls", "Grāmatvedība +", "Numo", "Projekta vadība"]

Noteikumi:
- "Instalācija" un "Projekta vadība" — vienmēr iekļauj
- "Pamatsistēma" — iekļauj ja minēta grāmatvedība, uzskaite, pamatsistēma, algas
- "Personāls" — iekļauj ja minēts personāls, HOP, pieteikumi, komandējumi, rīkojumi, darba laiks
- "Grāmatvedība +" — iekļauj ja minēti HOP rēķini vai rēķinu saskaņošana
- "Numo" — iekļauj ja minēta darba laika plānošana vai NUMO

Sarakstes vēsture:
{history_text}

Atbildi TIKAI JSON formātā, piemēram: ["Instalācija", "Pamatsistēma", "Projekta vadība"]"""

    model_cfg = MODELS[model_name]
    try:
        if model_cfg["provider"] == "mistral":
            client = load_mistral_client()
        else:
            client = load_gemini_client()
        if not client:
            return ["Instalācija", "Pamatsistēma", "Projekta vadība"]

        response = client.chat.completions.create(
            model=model_cfg["model"],
            messages=[{"role": "user", "content": prompt}],
        )
        import json
        text = response.choices[0].message.content.strip()
        if "```" in text:
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
        sections = json.loads(text)
        # Vienmēr pievieno Instalācija un Projekta vadība
        for always in ["Instalācija", "Projekta vadība"]:
            if always not in sections:
                sections.append(always)
        return sections
    except Exception:
        return ["Instalācija", "Pamatsistēma", "Projekta vadība"]


def generate_tame_excel(messages: list, model_name: str) -> tuple[bytes, str]:
    """Ģenerē ieviešanas tāmi Excel formātā, filtrējot blokus pēc sarakstes."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from copy import copy

    template_path = os.path.join(BASE_DIR, "assets", "tame_template.xlsx")
    if not os.path.exists(template_path):
        return None, "❌ Tāmes šablons nav atrasts (assets/tame_template.xlsx)"

    # Nosaka nepieciešamās sadaļas
    sections = determine_tame_sections(messages, model_name)

    # Iegūst klienta vārdu
    summary = summarize_chat_for_proposal(messages, model_name)
    klients = summary.get("klients", "Nav norādīts")
    timestamp = datetime.now().strftime("%d.%m.%Y")

    # Ielādē šablonu
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active

    # Atrod rindas ko dzēst (sadaļa nav iekļauta)
    rows_to_delete = []
    for row_idx in range(2, ws.max_row + 1):
        sadala = ws.cell(row=row_idx, column=1).value
        if sadala and sadala not in sections:
            rows_to_delete.append(row_idx)

    # Dzēš rindas no apakšas uz augšu
    for row_idx in reversed(rows_to_delete):
        ws.delete_rows(row_idx)

    # Iestata wrap_text un aprēķina rindu augstumu visām šūnām
    COL_WIDTHS = {1: 18, 2: 22, 3: 60, 4: 16, 5: 12, 6: 14, 7: 18}
    # Efektīvais rakstzīmju skaits uz rindu (mazāks nekā kolonnas platums — fonts un padding)
    EFFECTIVE_CHARS = {1: 15, 2: 18, 3: 48, 4: 13, 5: 10, 6: 11, 7: 15}
    LINE_HEIGHT = 15  # Excel punkti uz teksta rindu
    MIN_HEIGHT  = 32

    HEADER_ROW = 1   # galvenes rinda šablonā (pirms insert_rows)
    HEADER_HEIGHT = 22

    for row in ws.iter_rows():
        row_idx = row[0].row
        if row_idx == HEADER_ROW:
            # Galvenes rindai — fiksēts mazāks augstums
            for cell in row:
                if cell.value:
                    cell.alignment = Alignment(wrap_text=True, vertical="center",
                                               horizontal="center")
            ws.row_dimensions[row_idx].height = HEADER_HEIGHT
            continue

        max_lines = 1
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                eff_chars = EFFECTIVE_CHARS.get(cell.column, 18)
                text_lines = cell.value.split("\n")
                lines = 0
                for line in text_lines:
                    line = line.strip()
                    if line:
                        lines += max(1, -(-len(line) // eff_chars))
                    else:
                        lines += 1
                max_lines = max(max_lines, lines)
        row_height = max(MIN_HEIGHT, max_lines * LINE_HEIGHT)
        ws.row_dimensions[row_idx].height = row_height

    # Iestata kolonnu platumu
    for col_idx, width in COL_WIDTHS.items():
        col_letter = openpyxl.utils.get_column_letter(col_idx)
        ws.column_dimensions[col_letter].width = width

    # Pievieno klienta info virsrakstā (1. rinda pirms tabulas)
    ws.insert_rows(1)
    ws.insert_rows(1)
    ws["A1"] = f"Ieviešanas tāme — {klients}"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(wrap_text=False)
    ws.row_dimensions[1].height = 25
    ws["A2"] = f"Sagatavots: {timestamp}"
    ws["A2"].font = Font(italic=True)
    ws["A2"].alignment = Alignment(wrap_text=False)
    ws.row_dimensions[2].height = 20
    # Galvenes rinda pēc ievietošanas ir rinda 3 — fiksēts augstums
    ws.row_dimensions[3].height = 100

    # Saglabā atmiņā
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), sections


# ── Ziņojuma tipa klasifikators ───────────────────────────────────────────────

def _ir_dokumentu_jautajums(teksts: str, history: list) -> bool:
    """Nosaka vai ziņojums ir dokumentācijas jautājums (True → izmanto RAG)
    vai skaidra atbilde/komanda (False → RAG izlaists).
    Konservatīvs — šaubīgos gadījumos atgriež True (izmanto RAG).
    """
    t = teksts.strip().lower()
    vardi = t.split()
    pirmais = vardi[0] if vardi else ""

    # 0. Ja beidzas ar jautājuma zīmi — vienmēr izmanto RAG
    if teksts.rstrip().endswith("?"):
        return True

    # 1. Ļoti īss ziņojums (≤ 15 zīmes) — skaidra atbilde ("Jā", "Labi", "10 lietotāji")
    if len(t) <= 15:
        return False

    # 2. Sākas ar skaidru apstiprinājumu/noliegumu (tikai pirmais vārds)
    skaidras_atbildes = {"jā", "nē", "jap", "labi", "ok", "pareizi",
                         "tieši", "protams", "sapratu", "skaidrs", "paldies"}
    if pirmais in skaidras_atbildes:
        return False

    # 3. Tīra komanda bez jautājuma zīmes
    komandas = ["sagatavo ", "nosūti ", "gatavo ", "sūti "]
    if any(t.startswith(k) for k in komandas) and "?" not in teksts:
        return False

    # Visos pārējos gadījumos — izmanto RAG
    return True


# ── Automātiska rekvizītu iegūšana ───────────────────────────────────────────

_DETECT_IGNORET = {
    "Latvija", "Latvijā", "Latvijas", "Rīga", "Rīgā", "Rīgas",
    "Liepāja", "Liepājā", "Jelgava", "Jelgavā", "Jūrmala", "Jūrmalā",
    "Valmiera", "Daugavpils", "Jēkabpils", "Jūrmala", "Ventspils",
    "Eiropa", "Eiropā", "Eiropas", "Pasaule", "Pasaulē",
    "Paldies", "Lūdzu", "Jaunā", "Vecā", "Lielā", "Mazā",
}

def _detect_uznemums(teksts: str) -> str | None:
    """Mēģina atpazīt reģistrācijas numuru vai uzņēmuma nosaukumu tekstā."""

    # 1. Reģistrācijas numurs — 11 cipari
    reg_match = re.search(r'\b\d{11}\b', teksts)
    if reg_match:
        return reg_match.group(0)

    # 2. Uzņēmuma forma + nosaukums (SIA "X", AS X, IK X u.tml.)
    uzn_match = re.search(
        r'\b(SIA|AS|IK|ZS|BO|VSIA|VAS|PSIA|kooperatīvs?)\s+["\']?[\wĀāČčĒēĢģĪīĶķĻļŅņŠšŪūŽž][\wĀāČčĒēĢģĪīĶķĻļŅņŠšŪūŽž\s"\']{2,40}',
        teksts, re.IGNORECASE
    )
    if uzn_match:
        return uzn_match.group(0).strip().strip('"\'')

    # 3. Nosaukums pēdiņās (bez juridiskās formas)
    pedinju_match = re.search(r'["“„]([^"”“\n]{3,50})["”"]', teksts)
    if pedinju_match:
        return pedinju_match.group(1)

    # 4. 2–4 lielie vārdi pēc kārtas (nosaukums bez juridiskās formas)
    # Piemērs: "Latvijas Gāze", "Elme Messer", "Rimi Baltic"
    LV_UPPER = "A-ZĀČĒĢĪĶĻŅŠŪŽ"
    LV_LOWER = "a-zāčēģīķļņšūž"
    caps_matches = re.findall(
        rf'\b[{LV_UPPER}][{LV_LOWER}]{{2,}}(?:\s+[{LV_UPPER}][{LV_LOWER}]{{1,}}){{1,3}}\b',
        teksts
    )
    for candidate in caps_matches:
        words = candidate.split()
        if len(words) >= 2 and not any(w in _DETECT_IGNORET for w in words):
            return candidate

    return None


def _rekvizitu_konteksts(rek: dict) -> str:
    """Formatē iegūtos rekvizītus kā konteksta tekstu AI ziņojumam."""
    rindas = ["[Automātiski iegūti rekvizīti no LR Uzņēmumu reģistra — apstipriniet pie klienta]"]
    lauki = [
        ("Nosaukums",        rek.get("nosaukums", "")),
        ("Uzņēmuma tips",    rek.get("tips", "")),
        ("Reģ. nr.",         rek.get("reg_numurs", "")),
        ("Juridiskā adrese", rek.get("juridiska_adrese", "")),
        ("Reģistrēts",       rek.get("registrēts", "")),
    ]
    for label, val in lauki:
        if val:
            rindas.append(f"  {label}: {val}")

    # PVN statuss
    pvn_aktīvs = rek.get("pvn_aktīvs")
    if pvn_aktīvs is True:
        pvn_rinda = f"  PVN: {rek.get('pvn_numurs', '')} — aktīvs maksātājs (reģ. {rek.get('pvn_reģ', '')})"
        rindas.append(pvn_rinda)
    elif pvn_aktīvs is False:
        rindas.append("  PVN: nav reģistrēts kā PVN maksātājs")
    # pvn_aktīvs is None → API kļūda, nerakstām neko

    return "\n".join(rindas)


# ── Autentifikācija ───────────────────────────────────────────────────────────

def load_allowed_emails() -> set:
    """Ielādē atļauto e-pastu sarakstu.

    Avotu prioritāte:
    1. Streamlit Secrets: ALLOWED_EMAILS = "a@b.com,c@d.com"
    2. Fails allowed_emails.txt (viens e-pasts uz rindas, # — komentāri)
    """
    # 1. Streamlit Secrets
    try:
        raw = st.secrets.get("ALLOWED_EMAILS", "")
        if raw:
            return {e.strip().lower() for e in raw.split(",") if e.strip()}
    except Exception:
        pass

    # 2. Fails
    emails_file = os.path.join(BASE_DIR, "allowed_emails.txt")
    if os.path.exists(emails_file):
        emails = set()
        with open(emails_file, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#"):
                    emails.add(line.lower())
        return emails

    return set()


def show_login():
    """Rāda pieteikšanās lapu. Atgriež True, ja lietotājs autentificēts."""
    st.set_page_config(
        page_title="Pārdošanas aģents — pieteikšanās",
        page_icon="🔐",
        layout="centered",
    )

    # Centrētā karte
    col_l, col_c, col_r = st.columns([1, 2, 1])
    with col_c:
        logo_path = os.path.join(BASE_DIR, "assets", "logo.jpg")
        if os.path.exists(logo_path):
            st.image(logo_path, width=160)

        st.markdown("## 🔐 Pieteikšanās")
        st.markdown("Ievadi savu e-pasta adresi, lai piekļūtu sistēmai.")

        with st.form("login_form"):
            epasts = st.text_input("E-pasts", placeholder="vards.uzvards@uznemums.lv")
            submit = st.form_submit_button("Pieteikties", use_container_width=True)

        if submit:
            allowed = load_allowed_emails()
            if not allowed:
                st.error("⚠️ Nav konfigurēts atļauto lietotāju saraksts.")
                return False
            if epasts.strip().lower() in allowed:
                st.session_state.authenticated      = True
                st.session_state.authenticated_user = epasts.strip().lower()
                st.rerun()
            else:
                st.error("❌ Šis e-pasts nav reģistrēts. Sazinieties ar sistēmas administratoru.")

    return False


# ── Streamlit UI ──────────────────────────────────────────────────────────────

def main():
    st.set_page_config(
        page_title="pārdošanas aģents",
        page_icon="📚",
        layout="wide",
    )

    # ── Header ────────────────────────────────────────────────────────────────
    col_logo, col_title, col_img = st.columns([1, 3, 1])
    with col_logo:
        logo_path = os.path.join(BASE_DIR, "assets", "logo.jpg")
        if os.path.exists(logo_path):
            st.image(logo_path, use_container_width=True)
    with col_title:
        st.markdown("<h1 style='color: #003087;'>pārdošanas aģents</h1>", unsafe_allow_html=True)
    with col_img:
        img_path = os.path.join(BASE_DIR, "assets", "Horizon.jpg")
        if os.path.exists(img_path):
            st.image(img_path, use_container_width=True)

    # Indeksē un ielādē resursus
    auto_ingest()
    collection = load_collection()

    # ── Čats ──────────────────────────────────────────────────────────────────
    if "messages" not in st.session_state:
        st.session_state.messages = []

    def render_marker_buttons(markers: list, key_prefix: str):
        """Renderē darbību pogas atbilstoši AI atbildes marķieriem."""
        if not markers:
            return
        cols = st.columns(len(markers))
        for i, darbība in enumerate(markers):
            with cols[i]:
                if darbība == "tāme" and st.button(
                    "📊 Sagatavot ieviešanas tāmi", key=f"{key_prefix}_tāme",
                    use_container_width=True,
                ):
                    with st.spinner("Sagatavo tāmi..."):
                        excel_bytes, result = generate_tame_excel(
                            st.session_state.messages,
                            st.session_state.selected_model,
                        )
                    if excel_bytes:
                        sadaļas = ", ".join(result) if isinstance(result, list) else ""
                        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                        st.session_state.tame_bytes    = excel_bytes
                        st.session_state.tame_filename = f"horizon_tame_{ts}.xlsx"
                        st.session_state.tame_sadaļas  = sadaļas
                        ok, info = send_file_by_email(
                            excel_bytes, st.session_state.tame_filename,
                            subject=f"Horizon ieviešanas tāme — {datetime.now().strftime('%d.%m.%Y')}",
                            body=f"Pielikumā ieviešanas tāme.\nIekļautās sadaļas: {sadaļas}",
                            user_email=st.session_state.get("authenticated_user", ""),
                        )
                        if ok:
                            st.success(f"✅ Tāme sagatavota un nosūtīta uz: {info}")
                        else:
                            st.success(f"✅ Tāme sagatavota. Sadaļas: {sadaļas}")
                    else:
                        st.error(result)

                if darbība == "līgums" and st.button(
                    "📝 Sagatavot līgumu", key=f"{key_prefix}_līgums",
                    use_container_width=True,
                ):
                    with st.spinner("Sagatavo līgumu..."):
                        doc_bytes, rezultats = generate_ligums_docx(
                            st.session_state.messages,
                            st.session_state.selected_model,
                            st.session_state.get("klienta_rekviziti"),
                        )
                    if doc_bytes:
                        klienta_nos = rezultats.get("uznemuma_nosaukums", "") if isinstance(rezultats, dict) else ""
                        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                        st.session_state.ligums_bytes    = doc_bytes
                        st.session_state.ligums_mainīgie = rezultats
                        st.session_state.ligums_filename = f"horizon_ligums_{klienta_nos or ts}.docx"
                        ok, info = send_file_by_email(
                            doc_bytes, st.session_state.ligums_filename,
                            subject=f"Horizon līgums — {klienta_nos or datetime.now().strftime('%d.%m.%Y')}",
                            body=f"Pielikumā sagatavotais līgums{f' — {klienta_nos}' if klienta_nos else ''}.",
                            user_email=st.session_state.get("authenticated_user", ""),
                        )
                        if ok:
                            st.success(f"✅ Līgums sagatavots un nosūtīts uz: {info}")
                        else:
                            st.success("✅ Līgums sagatavots.")
                    else:
                        st.error(rezultats)

    for idx, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant":
                tables = extract_markdown_tables(msg["content"])
                if tables:
                    excel_bytes = tables_to_excel_bytes(tables)
                    ts = msg.get("timestamp", str(idx))
                    st.download_button(
                        label="📥 Lejupielādēt kā Excel",
                        data=excel_bytes,
                        file_name=f"horizon_aprekins_{ts}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"dl_{idx}",
                    )
                render_marker_buttons(msg.get("markers", []), key_prefix=f"msg_{idx}")
            if msg.get("sources"):
                with st.expander("📎 Avoti"):
                    for src in msg["sources"]:
                        st.text(f"• {src}")

    # ── Čata ievade ───────────────────────────────────────────────────────────────
    text_input = st.chat_input("Raksti vai ierunā jautājumu...")

    if "mic_counter" not in st.session_state:
        st.session_state.mic_counter = 0
    if "processed_audio_id" not in st.session_state:
        st.session_state.processed_audio_id = None

    audio_result = mic_button(reset_counter=st.session_state.mic_counter)

    question = None
    if text_input:
        question = text_input
    elif audio_result and audio_result.get("type") == "audio":
        # Unikāls ID — garums + beigu baiti (WebM sākums ir identisks visiem ierakstiem!)
        _d = audio_result["data"]
        audio_id = f"{len(_d)}-{_d[-40:]}"
        if audio_id != st.session_state.processed_audio_id:
            st.session_state.processed_audio_id = audio_id
            import base64 as _b64
            audio_bytes = _b64.b64decode(audio_result["data"])
            with st.spinner("Atpazīstu runu..."):
                question = transcribe_audio(audio_bytes, audio_result.get("mimeType", "audio/webm"))
            st.session_state.mic_counter += 1
            if not question:
                st.warning("⚠️ Neizdevās atpazīt runu — mēģini vēlreiz.")

    if question:
        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        with st.chat_message("assistant"):
            with st.spinner("Prātoju ..."):

                # Automātiska rekvizītu iegūšana — ja nav jau iegūti
                if not st.session_state.get("klienta_rekviziti"):
                    vaicajums = _detect_uznemums(question)
                    if vaicajums:
                        rek = fetch_data_gov_lv(vaicajums)
                        if "kļūda" not in rek:
                            st.session_state.klienta_rekviziti = rek

                rek_konteksts = ""
                if st.session_state.get("klienta_rekviziti"):
                    rek = st.session_state.klienta_rekviziti
                    rek_konteksts = "\n\n" + _rekvizitu_konteksts(rek)

                ir_jautajums = _ir_dokumentu_jautajums(question, st.session_state.messages)
                if ir_jautajums:
                    context, sources = retrieve_context(collection, question)
                else:
                    context, sources = "", []

                pilns_konteksts = (context or "") + rek_konteksts

                if not pilns_konteksts and ir_jautajums:
                    answer_raw = "Atbilstoši manām zināšanām, uz šo jautājumu nevarēšu sniegt precīzu atbildi. Jūsu jautājumu esmu piefiksējis un nosūtīšu menedžerim, kurš ar jums sazināsies tuvākajā laikā. [ACTION:NEZINA]"
                else:
                    answer_raw = ask_ai(
                        question, pilns_konteksts,
                        st.session_state.selected_model,
                        st.session_state.messages,
                        izmanto_rag=ir_jautajums or bool(rek_konteksts),
                    )

            answer, markers = parse_markers(answer_raw)

            st.markdown(answer)

            # Automātiska nosūtīšana ja AI nevar atbildēt
            if "nezina" in markers:
                user_email = st.session_state.get("authenticated_user", "")
                ok, sent_to = send_unanswered_question(question, user_email)
                if ok:
                    st.info(f"📧 Jautājums nosūtīts menedžerim uz: **{sent_to}**")
            tables = extract_markdown_tables(answer)
            if tables:
                excel_bytes = tables_to_excel_bytes(tables)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    label="📥 Lejupielādēt kā Excel",
                    data=excel_bytes,
                    file_name=f"horizon_aprekins_{timestamp}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            new_msg_idx = len(st.session_state.messages)
            render_marker_buttons(markers, key_prefix=f"msg_{new_msg_idx}")
            if sources:
                with st.expander("📎 Avoti"):
                    for src in sources:
                        st.text(f"• {src}")

        st.session_state.messages.append({
            "role":      "assistant",
            "content":   answer,
            "markers":   markers,
            "sources":   sources,
            "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S"),
        })

    # Modeļa noklusējums
    if "selected_model" not in st.session_state:
        st.session_state.selected_model = "🟣 Claude Sonnet"

    # Sānjosla
    with st.sidebar:

        # ── Lietotājs ─────────────────────────────────────────────────────────
        user = st.session_state.get("authenticated_user", "")
        st.caption(f"👤 {user}")
        if st.button("🚪 Izrakstīties"):
            if st.session_state.get("messages"):
                send_chat_by_email(
                    st.session_state.messages,
                    st.session_state.get("authenticated_user", ""),
                )
            st.session_state.messages           = []
            st.session_state.authenticated      = False
            st.session_state.authenticated_user = ""
            st.rerun()

        st.divider()

        # ── Klienta rekvizīti ─────────────────────────────────────────────────
        st.header("🏢 Klienta rekvizīti")
        reg_input = st.text_input(
            "Meklēt uzņēmumu:",
            placeholder="reģ. nr., nosaukums vai PVN nr.",
            key="firmas_meklet_input",
        )
        if st.button("🔍 Iegūt rekvizītus"):
            if not reg_input.strip():
                st.warning("⚠️ Ievadi reģistrācijas numuru, nosaukumu vai PVN numuru.")
            else:
                with st.spinner("Meklē firmas.lv..."):
                    rek = fetch_data_gov_lv(reg_input.strip())
                if "kļūda" in rek:
                    st.error(rek["kļūda"])
                else:
                    st.session_state.klienta_rekviziti = rek
                    st.success(f"✅ {rek.get('nosaukums', 'Atrasts!')}")
        if st.session_state.get("klienta_rekviziti"):
            rek = st.session_state.klienta_rekviziti
            with st.expander("📋 Rekvizīti", expanded=True):
                for atslega, nosaukums in [
                    ("nosaukums",        "Nosaukums"),
                    ("tips",             "Tips"),
                    ("reg_numurs",       "Reģ. nr."),
                    ("juridiska_adrese", "Juridiskā adrese"),
                    ("registrēts",       "Reģistrēts"),
                    ("sepa",             "SEPA"),
                ]:
                    val = rek.get(atslega, "")
                    if val:
                        st.caption(f"**{nosaukums}:** {val}")
                # PVN statuss
                pvn_aktīvs = rek.get("pvn_aktīvs")
                if pvn_aktīvs is True:
                    st.caption(f"**PVN:** ✅ {rek.get('pvn_numurs','')} (reģ. {rek.get('pvn_reģ','')})")
                elif pvn_aktīvs is False:
                    st.caption("**PVN:** ❌ nav reģistrēts")
                st.markdown(f"[🔗 UR.gov.lv]({rek.get('url', '')})")

        st.divider()

        # ── AI modelis ────────────────────────────────────────────────────────
        st.header("🤖 AI modelis")
        selected = st.selectbox(
            "Izvēlies modeli:",
            options=list(MODELS.keys()),
            index=list(MODELS.keys()).index(st.session_state.selected_model),
            key="model_selectbox",
            label_visibility="collapsed",
        )
        if selected != st.session_state.selected_model:
            st.session_state.selected_model = selected
            st.rerun()

        st.divider()

        # ── Sarakstes darbības ────────────────────────────────────────────────
        if st.button("📧 Nosūtīt sarunu uz e-pastu", use_container_width=True):
            ok, msg = send_chat_by_email(
                st.session_state.messages,
                st.session_state.get("authenticated_user", ""),
            )
            if ok:
                st.success(msg)
            else:
                st.error(msg)
        if st.button("🗑️ Notīrīt čatu", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
        if st.button("🔄 Pārindeksēt dokumentus", use_container_width=True):
            st.cache_resource.clear()
            st.rerun()


if __name__ == "__main__":
    if not st.session_state.get("authenticated"):
        show_login()
    else:
        main()
